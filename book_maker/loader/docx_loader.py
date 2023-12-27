import sys
from pathlib import Path
from bs4 import BeautifulSoup
from docx import Document
import io
from book_maker.utils import prompt_config_to_kwargs
import os
from .base_loader import BaseBookLoader
import boto3
import logging
s3 = boto3.client('s3')


class DocxBookLoader(BaseBookLoader):
    def __init__(
        self,
        docx_name,
        model,
        key,
        resume,
        language,
        model_api_base=None,
        is_test=False,
        test_num=5,
        prompt_config=None,
        single_translate=False,
        context_flag=False,
        temperature=1.0,
        bucket=None,
        upload_to_s3=False,
        language_key=None,

    ) -> None:
        self.docx_name = docx_name
        self.translate_model = model(
            key,
            language,
            api_base=model_api_base,
            temperature=temperature,
            **prompt_config_to_kwargs(prompt_config),
        )
        self.is_test = is_test
        self.p_to_save = []
        self.bilingual_result = []
        self.bilingual_temp_result = []
        self.test_num = test_num
        self.batch_size = 10
        self.single_translate = True
        self.bucket = bucket
        self.upload_to_s3 = False
        self.language = language
        self.language_key = language_key

        try:
            doc = Document(docx_name)
            self.origin_book = ""
            for paragraph in doc.paragraphs:
                self.origin_book += paragraph.text + "\n"

        except Exception as e:
            raise Exception("can not load file") from e

        self.resume = resume
        self.bin_path = f"{Path(docx_name).parent}/.{Path(docx_name).stem}.temp.bin"
        if self.resume:
            self.load_state()

    @staticmethod
    def _is_special_text(text):
        return text.isdigit() or text.isspace() or len(text) == 0

    def _make_new_book(self, book):
        pass

    def make_bilingual_book(self):
        index = 0
        p_to_save_len = len(self.p_to_save)
        logger = logging.getLogger()
        try:
            doc = Document(self.docx_name)
            p_tags = doc.paragraphs
            for p_tag in p_tags:
                batch_text = p_tag.text
                if self._is_special_text(batch_text):
                    continue
                if not self.resume or index >= p_to_save_len:
                    try:
                        temp = self.translate_model.translate(batch_text)
                    except Exception as e:
                        print(e)
                        raise Exception(
                            "Something is wrong when translate") from e
                    p_tag.text = temp
                    self.p_to_save.append(temp)
                    print(temp)
                    if not self.single_translate:
                        self.bilingual_result.append(batch_text)
                    self.bilingual_result.append(temp)
                index += 1
                if self.is_test and index > self.test_num:
                    break

            self.save_file(
                f"{Path(self.docx_name).parent}/tmp/{Path(self.docx_name).stem}.docx",
                doc,
            )

        except (KeyboardInterrupt, Exception) as e:
            print(e)
            print("you can resume it next time")
            self._save_progress()
            self._save_temp_book()
            sys.exit(0)

    def _save_temp_book(self):
        index = 0
        soup = BeautifulSoup(self.origin_book, 'html.parser')
        p_tags = soup.find_all('p')

        for p_tag in p_tags:
            batch_text = p_tag.get_text()
            self.bilingual_temp_result.append(batch_text)
            if self._is_special_text(batch_text):
                continue
            if index < len(self.p_to_save):
                self.bilingual_temp_result.append(self.p_to_save[index])
            index += 1

        self.save_file(
            f"{Path(self.docx_name).parent}/{Path(self.docx_name).stem}_bilingual_temp.docx",
            str(soup),
        )

    def _save_progress(self):
        try:
            with open(self.bin_path, "w", encoding="utf-8") as f:
                f.write("\n".join(self.p_to_save))
        except:
            raise Exception("can not save resume file")

    def load_state(self):
        try:
            with open(self.bin_path, encoding="utf-8") as f:
                self.p_to_save = f.read().splitlines()
        except Exception as e:
            raise Exception("can not load resume file") from e

    def save_file(self, book_path, doc):
        logger = logging.getLogger()
        if self.upload_to_s3:
            try:
                upload_path = '{}/{}.docx'.format(
                    self.language_key, Path(self.docx_name).stem)
                doc_content = io.StringIO()
                doc.save(doc_content)
                # Move the cursor to the beginning of the StringIO object
                doc_content.seek(0)
                # s3.put_object(Body=doc_content.read().decode('utf-8'),
                #               Bucket=self.bucket, Key=upload_path)
                logger.info(f"upload file to s3: {upload_path}")
                os.remove(book_path)
                logger.info(f"delete file: {book_path}")

            except (Exception) as e:
                print(e)
        else:
            try:
                doc.save(book_path)
                logger.info(f"save file: {book_path}")
            except (Exception) as e:
                print(e)
